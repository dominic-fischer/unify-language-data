import unicodedata
import string
import re
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
import pandas as pd
from bs4 import BeautifulSoup
import wikipediaapi
import requests
from io import StringIO
import os
from datetime import datetime
from pathlib import Path

class TextParser:
    def __init__(self):
        self.plaintext_lines = []
        self.headers = []
        self.table_snippets = []


    def build_markdown_plaintext(self):
        raise NotImplementedError


    def extract_headers(self):
        raise NotImplementedError


    def extract_table_snippets(self):
        raise NotImplementedError
    

    def save_output(self, output_path=None):
        final_text = self.insert_tables_into_plaintext()

        if output_path is None:
            folder = getattr(self, "output_folder", "output_files")
            Path(folder).mkdir(parents=True, exist_ok=True)
            filename = f"{self.title.replace(' ', '_')}_plaintext_structured_w_tables.txt"
            output_path = Path(folder) / filename

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_text)


    def normalize_and_map(self, text):
        """
        Normalize the input text by stripping accents and removing punctuation/whitespace.
        Also return a mapping from normalized characters back to original character indices.
        """
        norm = []
        mapping = []

        for i, orig_c in enumerate(text):
            for c in unicodedata.normalize('NFD', orig_c):
                if unicodedata.category(c) == 'Mn':
                    continue  # skip diacritic marks
                if c.isspace() or c in string.punctuation:
                    continue  # skip whitespace and punctuation
                norm.append(c.lower())
                mapping.append(i)

        return ''.join(norm), mapping


    def find_normalized_match(self, preceding, norm_text, mapping, original_text):
        """
        Try to find the normalized position of the preceding snippet in the normalized text.
        Returns the insertion point in the original text using the mapping, adjusted to skip
        punctuation, quotes, brackets, and whitespace after the match.
        """
        norm_snip = ''.join(
            c.lower() for c in self.strip_accents(preceding)
            if not c.isspace() and c not in string.punctuation
        )
        pos = norm_text.rfind(norm_snip)
        if pos == -1:
            return -1

        try:
            insert_pos = mapping[pos + len(norm_snip) - 1] + 1  # base insertion point

            # Skip over trailing punctuation, quotes, brackets, and whitespace
            while insert_pos < len(original_text) and original_text[insert_pos] in {
                '.', ',', '!', '?', ':', ';',
                '"', "'", '”', '’', '`', '“', '‘',
                ')', ']', '}', '›', '»', '-',
                ' '
            }:
                insert_pos += 1

            return insert_pos
        except IndexError:
            return -1


    def find_closest_header(self, text):
        """
        Return the last valid header from `self.headers` that appears before the end of `text`.
        Assumes `text` is the portion of the document up to the insertion point.
        """
        lines = text.splitlines()
        
        for line in reversed(lines):
            line = line.strip()
            if line.startswith('#'):  # likely a markdown header
                if line in self.headers:
                    return line

        return None  # fallback if no valid header found


    def find_level_header(self, header):
        """
        Find the level (number of '#') of the rightmost section in a breadcrumb-style header.
        For example, '## Grammar > ### Pronunciation' returns 3.
        """
        parts = header.split(' > ')
        if not parts:
            return 2
        last_part = parts[-1].strip()
        match = re.match(r'^(#+)', last_part)
        return len(match.group(1)) if match else 2


    def strip_accents(self, text):
        """
        Remove accents from characters in the input text using Unicode normalization.
        """
        return ''.join(
            c for c in unicodedata.normalize('NFD', text)
            if unicodedata.category(c) != 'Mn'
        )


    def insert_tables_into_plaintext(self):
        """
        Insert formatted tables into the plaintext at appropriate locations,
        based on normalized matches of preceding text and closest headers.
        """
        output = '\n'.join(self.plaintext_lines)
        norm_text, mapping = self.normalize_and_map(output)
        current_date_and_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        log_filename = f"debug_files/{self.title.replace(' ', '_')}_{current_date_and_time}_potential_errors_log.txt"

        for snippet in reversed(self.table_snippets):  # Reverse to avoid shifting insertion points
            preceding = snippet['preceding_text']
            table = snippet['table']
            title_str = snippet['title']
            fallback_header = "[No title available]"

            # Format the table
            table_str = table.to_markdown(index=False)
            table_block = f"<pre>```\n{table_str}\n```</pre>"

            # 1. Find normalized match location
            insert_pos = self.find_normalized_match(preceding, norm_text, mapping, output)
            if insert_pos == -1:
                print(f"❌ Could not find match for table {snippet['index']}: '{preceding}'")
                if not os.path.exists(log_filename):
                    open(log_filename, "w", encoding="utf-8").close()
                with open(log_filename, "a", encoding="utf-8") as f:
                    f.write("=== start ===\n")
                    f.write(f"table {snippet['index'] + 1} was not added to document.\n")
                    f.write("=== end ===\n")
                    f.write("\n")
                continue

            # 2. Get text up to that point and find closest header
            text_up_to = output[:insert_pos]
            closest_header = self.find_closest_header(text_up_to)
            #print(f"Closest header: '{closest_header}'")

            # Normalize both strings
            normalized_text_up_to = re.sub(r"\s+", "", text_up_to)
            normalized_preceding = re.sub(r"\s+", "", preceding)

            # Check and write to file if mismatch
            if not normalized_text_up_to.endswith(normalized_preceding):
                if not os.path.exists(log_filename):
                    open(log_filename, "w", encoding="utf-8").close()
                with open(log_filename, "a", encoding="utf-8") as f:
                    f.write("=== start ===\n")
                    f.write(f"Potential mismatch log for table {snippet['index'] + 1}\n")
                    f.write("=== text_up_to ===\n")
                    f.write(f"...'{normalized_text_up_to[-100:]}'\n")
                    f.write("=== preceding ===\n")
                    f.write(f"'{normalized_preceding}'\n")
                    f.write("=== end ===\n")
                    f.write("\n")

            print(f"Insertion point: ...'{text_up_to[-100:]}' > TABLE {snippet['index']}")
            print(f"Searching for closest header for table {snippet['index']} with preceding text: '{preceding}'")

            # 3. Determine header level and construct title
            level = self.find_level_header(closest_header) + 1
            title = f"{closest_header} > {'#' * level} {title_str}"
            if re.match(r'^(#+) Table \d+$', title):
                title += f": {fallback_header}"
            title_block = f"\n\n{title}\n\n{table_block}\n\n"

            # 4. Insert table into output
            output = output[:insert_pos] + title_block + output[insert_pos:]
            print(f"✅ Inserted table {snippet['index']} under '{closest_header}'\n")


        toc_string = "<pre>```\nTable of Contents\n"
        for h in self.headers:
            toc_string += h + "\n"
        toc_string += "```</pre>\n\n"

        output = toc_string + output

        return output.replace("\n\n\n", "\n\n")  # Clean up triple newlines


class DocxParser(TextParser):
    def __init__(self, file_path):
        super().__init__()
        self.output_folder = "docx_files"
        self.doc = Document(file_path)
        self.title = self.doc.core_properties.title
        self.plaintext_lines = [f"# {self.title}", ""]
        self.plaintext_lines += self.build_markdown_plaintext()
        self.headers = [f"# {self.title}"] + self.extract_headers()
        self.table_snippets = self.extract_table_snippets()


    def build_markdown_plaintext(self):
        """
        Convert DOCX content to markdown with breadcrumb-style headers.
        """
        lines = []
        current_heading_chain = []  # e.g., [(1, 'Verbs'), (2, 'Tenses and moods')]

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name

            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading", "").strip())

                    # Remove deeper levels from current chain
                    current_heading_chain = [
                        (lvl, t) for lvl, t in current_heading_chain if lvl < level
                    ]
                    current_heading_chain.append((level, text))

                    # Build breadcrumb header
                    breadcrumb = ' > '.join(
                        f"{'#' * (lvl + 1)} {title}" for lvl, title in current_heading_chain
                    )
                    lines.append(breadcrumb)
                    lines.append("")
                    continue
                except ValueError:
                    pass

            lines.append(text)
            lines.append("")

        return lines


    def extract_headers(self):
        """
        Extract headers from DOCX using breadcrumb-style chains.
        """
        headers = []
        current_heading_chain = []

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name

            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading", "").strip())

                    current_heading_chain = [
                        (lvl, t) for lvl, t in current_heading_chain if lvl < level
                    ]
                    current_heading_chain.append((level, text))

                    breadcrumb = ' > '.join(
                        f"{'#' * (lvl + 1)} {title}" for lvl, title in current_heading_chain
                    )
                    headers.append(breadcrumb)
                except ValueError:
                    continue

        return headers


    def extract_table_snippets(self):
        """
        Extract tables from the DOCX file with breadcrumb-style snippet context.
        The snippet is made of the last 10 words, including breadcrumb-style
        headers as part of the running word buffer.
        """
        ordered_elements = []
        for block in self.doc.element.body:
            if isinstance(block, CT_P):
                para = Paragraph(block, self.doc)
                ordered_elements.append(('paragraph', para))
            elif isinstance(block, CT_Tbl):
                table = Table(block, self.doc)
                ordered_elements.append(('table', table))

        table_snippets = []
        word_buffer = []
        current_heading_chain = []
        table_index = 0

        for item_type, item in ordered_elements:
            if item_type == 'paragraph':
                text = item.text.strip()
                if not text:
                    continue

                style = item.style.name
                if style.startswith("Heading"):
                    try:
                        level = int(style.replace("Heading", "").strip())
                        # Update the heading chain
                        current_heading_chain = [
                            (lvl, t) for lvl, t in current_heading_chain if lvl < level
                        ]
                        current_heading_chain.append((level, text))

                        # Add breadcrumb-style heading to buffer
                        breadcrumb = ' > '.join(
                            f"{'#' * (lvl + 1)} {t}" for lvl, t in current_heading_chain
                        )
                        word_buffer.extend(breadcrumb.split())
                        word_buffer = word_buffer[-50:]
                    except ValueError:
                        pass
                else:
                    # Add regular paragraph text to the buffer
                    words = text.split()
                    word_buffer.extend(words)
                    word_buffer = word_buffer[-50:]

            elif item_type == 'table':
                df = pd.DataFrame([
                    [cell.text.strip() for cell in row.cells]
                    for row in item.rows
                ])
                snippet = " ".join(word_buffer[-10:])

                table_snippets.append({
                    'index': table_index,
                    'preceding_text': snippet,
                    'table': df,
                    'title': f"Table {table_index + 1}"
                })
                table_index += 1

        return table_snippets


class WikipediaParser(TextParser):
    def __init__(self, page_name):
        super().__init__()
        self.output_folder = "wikipedia_files"
        self.page_name = page_name
        self.page = wikipediaapi.Wikipedia(
            user_agent='MyWikiParser/1.0',
            language='en'
        ).page(page_name)
        self.title = self.page.title
        self.plaintext_lines = [f"# {self.title}", "", self.page.summary.strip(), ""]
        self.plaintext_lines += self.build_markdown_plaintext()
        self.headers = [f"# {self.title}"] + self.extract_headers()
        self.table_snippets = self.extract_table_snippets()


    def build_markdown_plaintext(self, sections=None, level=0, parent_chain=None):
        """
        Build a markdown-formatted plaintext list with breadcrumb-style headers,
        and save section text under each full header.
        """
        if sections is None:
            sections = self.page.sections
        if parent_chain is None:
            parent_chain = []

        lines = []
        for section in sections:
            current_chain = parent_chain + [(level + 2, section.title)]

            # Build breadcrumb-style header
            breadcrumb = ' > '.join(f"{'#' * lvl} {title}" for lvl, title in current_chain)
            section_text = section.text.strip()

            lines.append(breadcrumb)
            lines.append("")
            lines.append(section_text)
            lines.append("")

            lines.extend(self.build_markdown_plaintext(section.sections, level + 1, current_chain))
        return lines


    def extract_headers(self, sections=None, level=0, parent_chain=None):
        """
        Recursively extract headers preserving each level's markdown and nesting.
        Example:
        ## Verbs
        ## Verbs > ### Tenses and moods
        ## Verbs > ### Tenses and moods > #### Imperative
        """
        if sections is None:
            sections = self.page.sections
        if parent_chain is None:
            parent_chain = []

        headers = []
        for section in sections:
            current_chain = parent_chain + [(level + 2, section.title)]  # level 0 → '##'

            # Build header string like:
            # ## Verbs > ### Tenses and moods > #### Imperative
            breadcrumb = ' > '.join(
                f"{'#' * lvl} {title}" for lvl, title in current_chain
            )
            headers.append(breadcrumb)

            headers.extend(
                self.extract_headers(section.sections, level + 1, current_chain)
            )
        return headers


    def extract_table_snippets(self):
        """
        Parse the HTML of the Wikipedia page and extract usable tables with snippet context.
        Uses breadcrumb-style section headers + surrounding context for robust matching.
        """
        url = f'https://en.wikipedia.org/wiki/{self.page_name.replace(" ", "_")}'
        html = requests.get(url).text
        soup = BeautifulSoup(html, 'html.parser')

        table_snippets = []
        word_buffer = []
        current_heading_chain = []  # e.g., [(2, 'Grammar'), (3, 'Pronunciation')]
        table_index = 0
        useful_index = 0

        for elem in soup.select('div.mw-parser-output > *'):
            # Update heading chain and add to word buffer
            if re.fullmatch(r'h[2-6]', elem.name):
                level = int(elem.name[1])
                title = elem.get_text(strip=True)

                current_heading_chain = [
                    (lvl, t) for (lvl, t) in current_heading_chain if lvl < level
                ]
                current_heading_chain.append((level, title))

                # Add breadcrumb to buffer
                breadcrumb = ' > '.join(
                    f"{'#' * lvl} {t}" for lvl, t in current_heading_chain
                )
                word_buffer.extend(breadcrumb.split())
                word_buffer = word_buffer[-50:]

            elif elem.name in ['p', 'ul', 'ol']:
                text = elem.get_text(" ", strip=True)
                words = text.split()
                word_buffer.extend(words)
                word_buffer = word_buffer[-50:]

            elif elem.name == 'table':
                try:
                    df = pd.read_html(StringIO(str(elem)))[0]
                    table_str = df.to_string(index=False).lower()

                    # Skip placeholder tables
                    placeholder_phrases = [
                        "this section needs expansion",
                        "you can help by adding to it",
                        "citation needed",
                        "unsourced material",
                        "learn how and when to remove this template message",
                        "is a stub",
                        "needs additional citations",
                        "may be challenged and removed",
                        "relevant discussion may be found",
                        "help improve this article",
                        "by adding citations to reliable sources"
                    ]
                    if any(phrase in table_str for phrase in placeholder_phrases):
                        print(f"Skipping table {table_index} due to placeholder content.")
                        print(f"Table content: {table_str}...\n")
                        table_index += 1
                        continue

                    caption = elem.caption.get_text(strip=True) if elem.caption else ""
                    caption = re.sub(r'\s*(\[\d+\])+\s*$', '', caption)

                    snippet = " ".join(word_buffer[-10:])

                    table_snippets.append({
                        'index': useful_index,
                        'preceding_text': snippet,
                        'table': df,
                        'title': f"Table {useful_index + 1}: {caption}" if caption else f"Table {useful_index + 1}"
                    })
                    useful_index += 1

                except Exception as e:
                    print(f"\tSkipping table {table_index} due to read_html error: {e}")
                table_index += 1

        print(f"Found {len(table_snippets)} tables with preceding text snippets.\n")
        for snippet in table_snippets:
            print(f"Table {snippet['index']} ({snippet['title']})")
            print(f"\tPreceding text: '{snippet['preceding_text']}'")
            print(f"\tTable shape: {snippet['table'].shape}\n")

        return table_snippets
